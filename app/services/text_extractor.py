"""
Text extraction from PDF and DOCX files
"""

import io
import fitz  # PyMuPDF
from docx import Document
import re
from typing import Union

def clean_text(text: str) -> str:
    """
    Clean and normalize extracted text
    
    Args:
        text: Raw text
    
    Returns:
        Cleaned text
    """
    # Remove extra whitespace
    text = re.sub(r'\s+', ' ', text)
    
    # Remove special characters but keep important ones
    text = re.sub(r'[^\w\s\-\+\#\.\,\@\(\)\:\;\/\&]', ' ', text)
    
    # Remove multiple spaces
    text = ' '.join(text.split())
    
    # Remove leading/trailing whitespace
    text = text.strip()
    
    return text

def extract_from_pdf(file_content: bytes) -> str:
    """
    Extract text from PDF using PyMuPDF
    
    Args:
        file_content: PDF file content as bytes
    
    Returns:
        Extracted text
    """
    try:
        # Open PDF from bytes
        pdf_document = fitz.open(stream=file_content, filetype="pdf")
        
        text_parts = []
        
        # Extract text from each page
        for page_num in range(pdf_document.page_count):
            page = pdf_document[page_num]
            page_text = page.get_text()
            
            if page_text.strip():
                text_parts.append(page_text)
        
        pdf_document.close()
        
        # Combine all pages
        full_text = "\n".join(text_parts)
        
        return clean_text(full_text)
    
    except Exception as e:
        print(f"❌ Error extracting PDF: {str(e)}")
        raise Exception(f"Failed to extract text from PDF: {str(e)}")

def extract_from_docx(file_content: bytes) -> str:
    """
    Extract text from DOCX file
    
    Args:
        file_content: DOCX file content as bytes
    
    Returns:
        Extracted text
    """
    try:
        # Open DOCX from bytes
        doc = Document(io.BytesIO(file_content))
        
        text_parts = []
        
        # Extract text from paragraphs
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text)
        
        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text)
                if row_text:
                    text_parts.append(" ".join(row_text))
        
        # Combine all text
        full_text = "\n".join(text_parts)
        
        return clean_text(full_text)
    
    except Exception as e:
        print(f"❌ Error extracting DOCX: {str(e)}")
        raise Exception(f"Failed to extract text from DOCX: {str(e)}")

def extract_text_from_file(file_content: bytes, filename: str) -> str:
    """
    Extract text from file based on extension
    
    Args:
        file_content: File content as bytes
        filename: Original filename
    
    Returns:
        Extracted text
    
    Raises:
        ValueError: If file type is not supported
    """
    filename_lower = filename.lower()
    
    if filename_lower.endswith('.pdf'):
        return extract_from_pdf(file_content)
    elif filename_lower.endswith('.docx'):
        return extract_from_docx(file_content)
    else:
        raise ValueError(f"Unsupported file type: {filename}. Only PDF and DOCX are supported.")